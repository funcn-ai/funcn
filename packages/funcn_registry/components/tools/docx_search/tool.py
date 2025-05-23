"""DOCX Search Tool for Word document processing and search."""

import aiofiles
import asyncio
import docx
import re
import xml.etree.ElementTree as ET
from datetime import datetime
from docx.document import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from pathlib import Path
from pydantic import BaseModel, Field
from typing import Any, Optional, Union
from zipfile import ZipFile


class DocumentSection(BaseModel):
    """Represents a section of text found in the document."""

    text: str = Field(..., description="The text content")
    paragraph_index: int = Field(..., description="Index of the paragraph in the document")
    section_type: str = Field(..., description="Type of section: paragraph, heading, table, etc.")
    style: str | None = Field(None, description="Style applied to this section")
    heading_level: int | None = Field(None, description="Heading level if this is a heading")
    table_info: dict[str, Any] | None = Field(None, description="Table information if this is from a table")
    match_positions: list[tuple[int, int]] = Field(default_factory=list, description="Start and end positions of matches")


class DOCXSearchResult(BaseModel):
    """Result of DOCX search operation."""

    success: bool = Field(..., description="Whether the search was successful")
    file_path: str = Field(..., description="Path to the DOCX file")
    total_matches: int = Field(..., description="Total number of matches found")
    matching_sections: list[DocumentSection] = Field(default_factory=list, description="Sections containing matches")
    document_metadata: dict[str, Any] = Field(default_factory=dict, description="Document metadata")
    error: str | None = Field(None, description="Error message if search failed")
    search_time: float = Field(..., description="Time taken to search in seconds")
    total_paragraphs: int = Field(0, description="Total number of paragraphs in document")
    total_tables: int = Field(0, description="Total number of tables in document")
    total_words: int = Field(0, description="Approximate total word count")


def validate_file_path(file_path: str) -> str:
    """Ensure the file exists and is a DOCX file."""
    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"File does not exist: {file_path}")
    if path.suffix.lower() not in ['.docx', '.docm']:
        raise ValueError(f"File must be a Word document (.docx or .docm): {file_path}")
    return str(path.absolute())


def extract_metadata(doc: Document) -> dict[str, Any]:
    """Extract metadata from the document."""
    metadata = {}

    try:
        core_props = doc.core_properties
        metadata.update({
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "keywords": core_props.keywords,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "last_modified_by": core_props.last_modified_by,
            "revision": core_props.revision,
            "category": core_props.category,
            "comments": core_props.comments
        })
    except Exception:
        pass

    return {k: v for k, v in metadata.items() if v is not None}


def search_in_text(
    text: str,
    search_text: str | None,
    regex_pattern: str | None,
    case_sensitive: bool,
    start_offset: int = 0
) -> list[tuple[int, int]]:
    """Search for matches in text and return positions."""
    if not text:
        return []

    matches = []

    if search_text:
        search_str = search_text if case_sensitive else search_text.lower()
        text_to_search = text if case_sensitive else text.lower()

        start = 0
        while True:
            pos = text_to_search.find(search_str, start)
            if pos == -1:
                break
            matches.append((start_offset + pos, start_offset + pos + len(search_str)))
            start = pos + 1

    if regex_pattern:
        flags = 0 if case_sensitive else re.IGNORECASE
        for match in re.finditer(regex_pattern, text, flags):
            matches.append((start_offset + match.start(), start_offset + match.end()))

    return matches


def get_heading_level(paragraph: Paragraph) -> int | None:
    """Get heading level from paragraph style."""
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.split()[-1])
        except (ValueError, IndexError):
            pass
    return None


def extract_table_text(table: Table) -> list[list[dict[str, Any]]]:
    """Extract text from a table with structure information."""
    table_data = []

    for row_idx, row in enumerate(table.rows):
        row_text = []
        for col_idx, cell in enumerate(row.cells):
            cell_text = ' '.join(paragraph.text for paragraph in cell.paragraphs)
            row_text.append({
                "row": row_idx,
                "col": col_idx,
                "text": cell_text
            })
        table_data.append(row_text)

    return table_data


def should_include_section(
    paragraph: Paragraph,
    style_filter: str | None,
    extract_headings: bool,
    heading_level: int | None
) -> bool:
    """Check if a section should be included based on filters."""
    if style_filter and paragraph.style.name != style_filter:
        return False

    if extract_headings:
        paragraph_heading_level = get_heading_level(paragraph)
        if paragraph_heading_level is None:
            return False
        if heading_level is not None and paragraph_heading_level != heading_level:
            return False

    return True


async def process_docx_document(
    file_path: str,
    search_text: str | None = None,
    regex_pattern: str | None = None,
    case_sensitive: bool = False,
    include_tables: bool = True,
    extract_headings: bool = False,
    heading_level: int | None = None,
    style_filter: str | None = None,
    max_context_chars: int = 200,
    extract_metadata_flag: bool = True,
    include_comments: bool = False
) -> DOCXSearchResult:
    """Process the DOCX document and search for content.

    Args:
        file_path: Path to the DOCX file
        search_text: Text to search for (case-insensitive by default)
        regex_pattern: Regular expression pattern to search for
        case_sensitive: Whether search should be case sensitive
        include_tables: Whether to search within tables
        extract_headings: Extract all headings from the document
        heading_level: Specific heading level to extract (1-9)
        style_filter: Only search in sections with this style
        max_context_chars: Maximum characters of context around matches
        extract_metadata_flag: Whether to extract document metadata
        include_comments: Whether to include document comments

    Returns:
        DOCXSearchResult with search results
    """
    start_time = asyncio.get_event_loop().time()

    try:
        # Validate file path
        file_path = validate_file_path(file_path)

        # Load the document
        doc = docx.Document(file_path)

        # Extract metadata if requested
        metadata = extract_metadata(doc) if extract_metadata_flag else {}

        matching_sections = []
        total_words = 0
        paragraph_index = 0

        # Process paragraphs
        for paragraph in doc.paragraphs:
            if not should_include_section(paragraph, style_filter, extract_headings, heading_level):
                continue

            text = paragraph.text
            if not text.strip():
                continue

            total_words += len(text.split())

            # Search for matches or extract if it's a heading
            matches = search_in_text(text, search_text, regex_pattern, case_sensitive)
            paragraph_heading_level = get_heading_level(paragraph)

            if matches or (extract_headings and paragraph_heading_level is not None):
                section = DocumentSection(
                    text=text,
                    paragraph_index=paragraph_index,
                    section_type="heading" if paragraph_heading_level else "paragraph",
                    style=paragraph.style.name,
                    heading_level=paragraph_heading_level,
                    table_info=None,
                    match_positions=matches
                )
                matching_sections.append(section)

            paragraph_index += 1

        # Process tables if requested
        table_count = 0
        if include_tables:
            for table_idx, table in enumerate(doc.tables):
                table_count += 1
                table_data = extract_table_text(table)

                # Search in table cells
                for row_data in table_data:
                    for cell_data in row_data:
                        cell_text = cell_data["text"]
                        matches = search_in_text(cell_text, search_text, regex_pattern, case_sensitive)

                        if matches:
                            section = DocumentSection(
                                text=cell_text,
                                paragraph_index=paragraph_index,
                                section_type="table",
                                style=None,
                                heading_level=None,
                                table_info={
                                    "table_index": table_idx,
                                    "row": cell_data["row"],
                                    "col": cell_data["col"]
                                },
                                match_positions=matches
                            )
                            matching_sections.append(section)

                paragraph_index += 1

        search_time = asyncio.get_event_loop().time() - start_time

        return DOCXSearchResult(
            success=True,
            file_path=file_path,
            total_matches=sum(len(s.match_positions) for s in matching_sections),
            matching_sections=matching_sections,
            document_metadata=metadata,
            error=None,
            search_time=search_time,
            total_paragraphs=len(doc.paragraphs),
            total_tables=table_count,
            total_words=total_words
        )

    except Exception as e:
        search_time = asyncio.get_event_loop().time() - start_time
        return DOCXSearchResult(
            success=False,
            file_path=file_path,
            total_matches=0,
            matching_sections=[],
            document_metadata={},
            error=str(e),
            search_time=search_time,
            total_paragraphs=0,
            total_tables=0,
            total_words=0
        )


# Convenience functions
async def search_docx(
    file_path: str,
    search_text: str,
    case_sensitive: bool = False,
    include_tables: bool = True
) -> DOCXSearchResult:
    """Search for text in a DOCX file.

    Args:
        file_path: Path to the DOCX file
        search_text: Text to search for
        case_sensitive: Whether search should be case sensitive
        include_tables: Whether to search in tables

    Returns:
        DOCXSearchResult with matching sections
    """
    return await process_docx_document(
        file_path=file_path,
        search_text=search_text,
        case_sensitive=case_sensitive,
        include_tables=include_tables
    )


async def extract_docx_headings(
    file_path: str,
    heading_level: int | None = None
) -> DOCXSearchResult:
    """Extract all headings from a DOCX file.

    Args:
        file_path: Path to the DOCX file
        heading_level: Specific heading level to extract (1-9)

    Returns:
        DOCXSearchResult with heading sections
    """
    return await process_docx_document(
        file_path=file_path,
        extract_headings=True,
        heading_level=heading_level
    )


async def search_docx_with_regex(
    file_path: str,
    pattern: str,
    include_tables: bool = True
) -> DOCXSearchResult:
    """Search DOCX file using regular expression.

    Args:
        file_path: Path to the DOCX file
        pattern: Regular expression pattern
        include_tables: Whether to search in tables

    Returns:
        DOCXSearchResult with matching sections
    """
    return await process_docx_document(
        file_path=file_path,
        regex_pattern=pattern,
        include_tables=include_tables
    )
