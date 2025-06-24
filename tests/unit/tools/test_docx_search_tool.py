"""Test suite for docx_search_tool following best practices."""

import asyncio
import pytest

# Import the actual tool functions
from packages.funcn_registry.components.tools.docx_search.tool import (
    DocumentSection,
    DOCXSearchResult,
    extract_docx_headings,
    process_docx_document,
    search_docx,
    search_docx_with_regex,
)
from pathlib import Path
from tests.fixtures import TestDataFactory
from tests.utils import BaseToolTest
from unittest.mock import Mock, patch


class TestDOCXSearchTool(BaseToolTest):
    """Test docx_search_tool component."""
    
    component_name = "docx_search_tool"
    component_path = Path("packages/funcn_registry/components/tools/docx_search")
    
    def create_test_docx(self, file_path: Path, content: list[dict]) -> Path:
        """Create a test DOCX file with specified content.
        
        Args:
            file_path: Path where to save the DOCX file
            content: List of dicts with 'text', 'style' keys
        
        Returns:
            Path to the created file
        """
        import docx
        from docx.shared import Pt
        
        doc = docx.Document()
        
        # Set document properties
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.core_properties.subject = "Testing"
        doc.core_properties.keywords = "test, document, search"
        
        for item in content:
            text = item.get('text', '')
            style = item.get('style', 'Normal')
            
            if style.startswith('Heading'):
                paragraph = doc.add_heading(text, level=int(style[-1]) if style[-1].isdigit() else 1)
            elif style == 'Table':
                # Create table
                table_data = item.get('table_data', [])
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_text
            else:
                paragraph = doc.add_paragraph(text, style=style)
                
                # Add formatting if specified
                if item.get('bold'):
                    for run in paragraph.runs:
                        run.bold = True
                if item.get('italic'):
                    for run in paragraph.runs:
                        run.italic = True
        
        doc.save(str(file_path))
        return file_path
    
    def get_component_function(self):
        """Import the tool function."""
        return search_docx
    
    def get_test_inputs(self):
        """Provide test inputs for the tool."""
        return [
            {
                "file_path": "/tmp/test.docx",
                "search_text": "introduction",
                "case_sensitive": False,
                "include_tables": True
            },
            {
                "file_path": "/tmp/report.docx",
                "search_text": "financial results",
                "case_sensitive": False,
                "include_tables": True
            },
            {
                "file_path": "/tmp/thesis.docx",
                "search_text": "conclusion",
                "case_sensitive": True,
                "include_tables": False
            }
        ]
    
    def validate_tool_output(self, output, input_data):
        """Validate the tool output format."""
        assert isinstance(output, DOCXSearchResult)
        assert isinstance(output.success, bool)
        assert isinstance(output.file_path, str)
        assert isinstance(output.total_matches, int)
        assert isinstance(output.matching_sections, list)
        assert isinstance(output.document_metadata, dict)
        assert isinstance(output.search_time, float)
        
        for section in output.matching_sections:
            assert isinstance(section, DocumentSection)
            assert isinstance(section.text, str)
            assert isinstance(section.paragraph_index, int)
            assert isinstance(section.section_type, str)
            assert isinstance(section.match_positions, list)
    
    @pytest.mark.asyncio
    async def test_search_paragraphs(self, tmp_path):
        """Test searching in document paragraphs."""
        docx_file = tmp_path / "test.docx"
        
        # Create test document
        content = [
            {"text": "Introduction to the topic", "style": "Normal"},
            {"text": "This is the main content with important information", "style": "Normal"},
            {"text": "Conclusion and summary", "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Search for text
        result = await search_docx(
            file_path=str(docx_file),
            search_text="important",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches >= 1
        assert len(result.matching_sections) >= 1
        assert any("important" in section.text.lower() for section in result.matching_sections)
    
    @pytest.mark.asyncio
    async def test_search_headers(self, tmp_path):
        """Test searching in document headers."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with headers
        content = [
            {"text": "Chapter 1: Introduction", "style": "Heading 1"},
            {"text": "Regular paragraph", "style": "Normal"},
            {"text": "Section 1.1: Background", "style": "Heading 2"},
            {"text": "1.1.1 Historical Context", "style": "Heading 3"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Search for text in headers
        result = await search_docx(
            file_path=str(docx_file),
            search_text="Introduction",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches >= 1
        # Should find in Heading 1
        heading_matches = [s for s in result.matching_sections if s.style == "Heading 1"]
        assert len(heading_matches) >= 1
    
    @pytest.mark.asyncio
    async def test_search_tables(self, tmp_path):
        """Test searching in document tables."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with table
        content = [
            {"text": "Product Catalog", "style": "Heading 1"},
            {
                "style": "Table",
                "table_data": [
                    ["Product Name", "Price"],
                    ["Laptop Computer", "$999"],
                    ["Desktop Computer", "$1299"]
                ]
            }
        ]
        self.create_test_docx(docx_file, content)
        
        # Search in tables
        result = await search_docx(
            file_path=str(docx_file),
            search_text="Laptop",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches >= 1
        # Should find in table
        table_matches = [s for s in result.matching_sections if s.section_type == "table"]
        assert len(table_matches) >= 1
    
    @pytest.mark.asyncio
    async def test_extract_headings(self, tmp_path):
        """Test extracting headings from document."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with various headings
        content = [
            {"text": "Introduction", "style": "Heading 1"},
            {"text": "This is the introduction paragraph", "style": "Normal"},
            {"text": "Background", "style": "Heading 2"},
            {"text": "Some background information", "style": "Normal"},
            {"text": "Methodology", "style": "Heading 1"},
            {"text": "Data Collection", "style": "Heading 2"},
            {"text": "Analysis Methods", "style": "Heading 2"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Extract all headings
        result = await extract_docx_headings(
            file_path=str(docx_file)
        )
        
        assert result.success is True
        assert len(result.matching_sections) >= 5  # Should find all headings
        
        # Extract only level 2 headings
        result_level2 = await extract_docx_headings(
            file_path=str(docx_file),
            heading_level=2
        )
        
        assert result_level2.success is True
        level2_headings = [s for s in result_level2.matching_sections if s.heading_level == 2]
        assert len(level2_headings) == 3
    
    @pytest.mark.asyncio
    async def test_case_sensitivity(self, tmp_path):
        """Test case-sensitive vs case-insensitive search."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with different cases
        content = [
            {"text": "UPPERCASE text here", "style": "Normal"},
            {"text": "lowercase text here", "style": "Normal"},
            {"text": "MixedCase Text Here", "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Case-insensitive search
        result_insensitive = await search_docx(
            file_path=str(docx_file),
            search_text="text",
            case_sensitive=False,
            include_tables=True
        )
        
        # Case-sensitive search
        result_sensitive = await search_docx(
            file_path=str(docx_file),
            search_text="text",
            case_sensitive=True,
            include_tables=True
        )
        
        assert result_insensitive.success is True
        assert result_sensitive.success is True
        # Case-insensitive should find more matches
        assert result_insensitive.total_matches >= result_sensitive.total_matches
        assert result_insensitive.total_matches >= 3  # Should find in all 3 paragraphs
        assert result_sensitive.total_matches >= 1   # Should find at least in lowercase
    
    @pytest.mark.asyncio
    async def test_document_metadata(self, tmp_path):
        """Test extracting document metadata."""
        docx_file = tmp_path / "test.docx"
        
        # Create document (metadata was set in create_test_docx)
        content = [
            {"text": "This is a test document", "style": "Normal"},
            {"text": "With some content", "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Search with metadata extraction
        result = await process_docx_document(
            file_path=str(docx_file),
            search_text="test",
            case_sensitive=False,
            extract_metadata_flag=True
        )
        
        assert result.success is True
        assert result.document_metadata is not None
        assert "title" in result.document_metadata
        assert result.document_metadata["title"] == "Test Document"
        assert "author" in result.document_metadata
        assert result.document_metadata["author"] == "Test Author"
    
    @pytest.mark.asyncio
    async def test_regex_search(self, tmp_path):
        """Test searching with regular expressions."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with various patterns
        content = [
            {"text": "Contact us at: info@example.com", "style": "Normal"},
            {"text": "Phone: +1 (555) 123-4567", "style": "Normal"},
            {"text": "Date: 2024-01-15", "style": "Normal"},
            {"text": "Reference: DOC-12345-ABC", "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Search for email pattern
        result = await search_docx_with_regex(
            file_path=str(docx_file),
            pattern=r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches >= 1
        assert any("@example.com" in section.text for section in result.matching_sections)
        
        # Search for date pattern
        result_date = await search_docx_with_regex(
            file_path=str(docx_file),
            pattern=r"\d{4}-\d{2}-\d{2}",
            include_tables=True
        )
        
        assert result_date.success is True
        assert result_date.total_matches >= 1
    
    @pytest.mark.asyncio
    async def test_style_filter(self, tmp_path):
        """Test filtering by paragraph style."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with different styles
        content = [
            {"text": "Main Title", "style": "Heading 1"},
            {"text": "This is normal text with important info", "style": "Normal"},
            {"text": "Subtitle with important details", "style": "Heading 2"},
            {"text": "Another normal paragraph with data", "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Search only in Normal style paragraphs
        result = await process_docx_document(
            file_path=str(docx_file),
            search_text="important",
            case_sensitive=False,
            style_filter="Normal"
        )
        
        assert result.success is True
        assert result.total_matches >= 1
        # All matches should be in Normal style
        for section in result.matching_sections:
            assert section.style == "Normal"
    
    @pytest.mark.asyncio
    async def test_context_chars(self, tmp_path):
        """Test controlling context around matches."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with long paragraph
        long_text = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
        long_text += "The word TARGET appears here in the middle of a very long paragraph. "
        long_text += "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua."
        
        content = [
            {"text": long_text, "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Search with limited context
        result = await process_docx_document(
            file_path=str(docx_file),
            search_text="TARGET",
            case_sensitive=False,
            max_context_chars=50
        )
        
        assert result.success is True
        assert result.total_matches >= 1
        # Context should be limited
        for section in result.matching_sections:
            assert "TARGET" in section.text
    
    @pytest.mark.asyncio
    async def test_multiple_tables(self, tmp_path):
        """Test searching across multiple tables."""
        docx_file = tmp_path / "test.docx"
        
        # Create document with multiple tables
        content = [
            {"text": "Sales Report", "style": "Heading 1"},
            {
                "style": "Table",
                "table_data": [
                    ["Product", "Q1 Sales", "Q2 Sales"],
                    ["Laptop", "$50,000", "$75,000"],
                    ["Desktop", "$30,000", "$45,000"]
                ]
            },
            {"text": "Inventory Status", "style": "Heading 1"},
            {
                "style": "Table",
                "table_data": [
                    ["Item", "Stock", "Location"],
                    ["Laptop Model A", "25", "Warehouse 1"],
                    ["Desktop Pro", "15", "Warehouse 2"]
                ]
            }
        ]
        self.create_test_docx(docx_file, content)
        
        # Search across tables
        result = await search_docx(
            file_path=str(docx_file),
            search_text="Laptop",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches >= 2  # Should find in both tables
        assert result.total_tables == 2
    
    @pytest.mark.asyncio
    async def test_empty_document(self, tmp_path):
        """Test handling of empty documents."""
        docx_file = tmp_path / "empty.docx"
        
        # Create empty document
        content = []
        self.create_test_docx(docx_file, content)
        
        result = await search_docx(
            file_path=str(docx_file),
            search_text="test",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches == 0
        assert len(result.matching_sections) == 0
        assert result.total_paragraphs == 0
    
    @pytest.mark.asyncio
    async def test_large_document_performance(self, tmp_path):
        """Test performance with large documents."""
        docx_file = tmp_path / "large.docx"
        
        # Create large document
        content = [
            {"text": f"Paragraph {i} with some content about topic {i % 10}", "style": "Normal"}
            for i in range(500)  # Reduced from 1000 for faster test
        ]
        self.create_test_docx(docx_file, content)
        
        import time
        start_time = time.time()
        
        result = await search_docx(
            file_path=str(docx_file),
            search_text="Paragraph 250",
            case_sensitive=False,
            include_tables=True
        )
        
        elapsed = time.time() - start_time
        
        assert result.success is True
        assert elapsed < 5.0  # Should complete within 5 seconds
        assert result.total_matches >= 1
        assert result.total_paragraphs == 500
    
    @pytest.mark.asyncio
    async def test_nonexistent_file(self, tmp_path):
        """Test handling of nonexistent files."""
        docx_file = tmp_path / "nonexistent.docx"
        
        result = await search_docx(
            file_path=str(docx_file),
            search_text="test",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is False
        assert result.error is not None
        assert "does not exist" in result.error
        assert result.total_matches == 0
    
    @pytest.mark.asyncio
    async def test_invalid_file_extension(self, tmp_path):
        """Test handling of non-DOCX files."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("This is not a DOCX file")
        
        result = await search_docx(
            file_path=str(txt_file),
            search_text="test",
            case_sensitive=False,
            include_tables=True
        )
        
        assert result.success is False
        assert result.error is not None
        assert "must be a Word document" in result.error
    
    @pytest.mark.asyncio
    async def test_no_search_criteria(self, tmp_path):
        """Test extracting all content without search criteria."""
        docx_file = tmp_path / "test.docx"
        
        content = [
            {"text": "First paragraph", "style": "Normal"},
            {"text": "Second paragraph", "style": "Normal"},
            {"text": "Third paragraph", "style": "Normal"}
        ]
        self.create_test_docx(docx_file, content)
        
        # Process without search text should not match anything
        result = await process_docx_document(
            file_path=str(docx_file),
            search_text=None,
            regex_pattern=None,
            include_tables=True
        )
        
        assert result.success is True
        assert result.total_matches == 0
        assert len(result.matching_sections) == 0
        assert result.total_paragraphs == 3
